import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_markdown_to_docx(input_md, output_docx):
    if not os.path.exists(input_md):
        print(f"Error: {input_md} not found.")
        return

    doc = Document()
    
    # Read Markdown content
    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Define Academic Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        
        # Chapter Headings (Giant center titles to hit page count)
        if line.startswith('# '):
            if any(key in line for key in ["Chapter", "Report", "CERTIFICATE", "ABSTRACT", "ACKNOWLEDGEMENT", "SUMMARY", "TABLE OF CONTENTS", "LIST OF"]):
                doc.add_page_break()
            
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Force large size for chapters
            run = h.runs[0]
            run.font.size = Pt(26)
            run.bold = True

        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
            run = h.runs[0]
            run.font.size = Pt(18)
            
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Horizontal Rules (Page break)
        elif line == '---':
            doc.add_page_break()
            
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.line_spacing = 1.15

        # Tables
        elif '|' in line:
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells: continue
            
            p = doc.add_paragraph("\t".join(cells))
            p.paragraph_format.line_spacing = 1.5
            p.style = 'Quote' # Uses a distinct style for tables

        # Placeholder boxes
        elif '[FIGURE' in line or '[PLACEHOLDER' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"\n\n--- {line} ---\n\n")
            run.bold = True
            run.font.size = Pt(12)
            run.italic = True

        # Body Text
        elif line:
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph(clean_line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Set Justify Alignment
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(12) # Academic paragraph gap
            
            if line.startswith('**') and line.endswith('**'):
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)

    doc.save(output_docx)
    print(f"Successfully converted {input_md} to {output_docx}")

if __name__ == "__main__":
    # Path to the report I just generated
    input_file = r"Academic_Project_Report.md"
    output_file = r"PawDevX_Industrial_Training_Report.docx"
    
    convert_markdown_to_docx(input_file, output_file)
